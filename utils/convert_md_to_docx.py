"""Convert Markdown resume to DOCX (Word document) format."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys


def markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX format.
    
    Args:
        md_file: Path to input Markdown file
        docx_file: Path to output DOCX file
    """
    # Read the markdown file
    with open(md_file, 'r') as f:
        content = f.read()

    # Create a new Document
    doc = Document()

    # Process line by line
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            # Heading 1
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip() == '':
            # Empty line / paragraph break
            doc.add_paragraph()
        else:
            # Regular paragraph
            if line.strip():
                doc.add_paragraph(line)

    # Save the document
    doc.save(docx_file)
    print(f"✓ Converted {md_file} → {docx_file}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python convert_md_to_docx.py input.md output.docx")
        sys.exit(1)
    markdown_to_docx(sys.argv[1], sys.argv[2])
